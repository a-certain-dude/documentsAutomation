from docx import Document

# Create the Word document
doc = Document()
doc.add_heading('ST. FRANCIS OF ASSISI CATHOLIC JUNIOR HIGH SCHOOL', level=1)
doc.add_heading('SCHOOL BASED ASSESSMENT (SBA) FORM', level=2)
doc.add_paragraph('TERM: ..............     FORM: 1A BOYS     SUBJECT TEACHER: ..............................')
doc.add_paragraph('SUBJECT: .....................................................     YEAR: .......................')

# Define table headers
headers = [
    "NO.", "NAME", "CAT 1\n10", "GROUP WORK\n20", "CAT 2\n20",
    "SUB TOTAL CATS\n50", "END OF TERM EXAMS\n100", "50% OF END OF TERM EXAMS",
    "TOTAL CATS + EXAMS", "GRADE"
]

# Table data from the image
table_data = [
    [1, "ADAMU SALIFU", 0, 0, 0],
    [2, "AGBANYO CLETUS", 10, 2, 7],
    [3, "AGBOGAH CARL", 0, 6, 0],
    [4, "AGBOTTOME JOHN", 0, 6, 8],
    [5, "AHADZI CLEMENT", 10, 6, 10],
    [6, "AHIABU GLADSTONE", 0, 0, 0],
    [7, "AKRAH GODSWAY", 6, 0, 5],
    [8, "ALORKPA MOSES JOHN", 10, 19, 15],
    [9, "ATSITSOGBUI KELLY", 10, 7, 15],
    [10, "ATSU BRIGHT", 10, 7, 14],
    [11, "BATAKA DELADEM", 10, 5, 7],
    [12, "BESSAH WILLIAM", 5, 5, 6],
    [13, "DAKETSE DANIEL", 7, 2, 8],
    [14, "DEKU COURAGE", 10, 12, 16],
    [15, "DOGBA MOSES", 10, 4, 2],
    [16, "DOGBE EMMANUEL", 10, 2, 10],
    [17, "DZIWORNU PIUS", 10, 10, 11],
    [18, "GADOR MAXWELL", 10, 10, 14],
    [19, "GELI WISDOM", 0, 0, 6],
    [20, "HIAMEY MASSEL", 10, 12, 13],
    [21, "KPODO SUCCESS DEBRAH", 10, 12, 13],
    [22, "MORTEY FRANCIS", 10, 6, 7],
    [23, "NYAMADI MANVEL", 10, 8, 7],
    [24, "TOMFEYA NOEL", 10, 2, 3],
    [25, "TSALIM STEPHEN", 10, 3, 4],
    [26, "YUSIF HABIB", 10, 4, 0]
]

# Create the table
table = doc.add_table(rows=1, cols=len(headers))
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    hdr_cells[i].text = header

# Fill the table
for row in table_data:
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = str(val)
    for _ in range(5):  # Add 5 blank cells for the rest of the columns
        row_cells[len(row)].text = ""

# Save the file
doc.save("SBA_Form_1A_Boys.docx")
print("SBA form created successfully.")
