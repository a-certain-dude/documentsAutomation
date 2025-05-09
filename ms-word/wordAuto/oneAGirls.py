from docx import Document

# Create the Word document
doc = Document()
doc.add_heading('ST. FRANCIS OF ASSISI CATHOLIC JUNIOR HIGH SCHOOL', level=1)
doc.add_heading('SCHOOL BASED ASSESSMENT (SBA) FORM', level=2)
doc.add_paragraph('TERM: ..............     FORM: 1A BOYS     SUBJECT TEACHER: ..............................')
doc.add_paragraph('SUBJECT: .....................................................     YEAR: .......................')

# Define table headers
headers = [
    "NO.", "NAME", "CAT 1\n10", "GROUP WORK\n20", "CAT 2\n20",
    "SUB TOTAL CATS\n50", "END OF TERM EXAMS\n100", "50% OF END OF TERM EXAMS",
    "TOTAL CATS + EXAMS", "GRADE"
]

# Table data from the image
table_data = [
    [1, "ADABADZE SCHORLASTICA", 10, 6, 16],
    [2, "ADABRA RUTH", 10, 5, 10],
    [3, "ADAMU BARAKATU", 0, 9, 16],
    [4, "AGBADZI HANNAH", 10, 8, 12],
    [5, "AGBLEY LAWRENCIA", 4, 9, 16],
    [6, "AGLAGO RACHAEL", 6,0, 11],
    [7, "AHIABLE ESTHER", 10, 14, 12],
    [8, "AKAKRO JUDITH", 7, 6, 17],
    [9, "AKORLI PRINCELLA", 10, 12, 11],
    [10, "ANATOR YORM RUBY", 10, 12, 17],
    [11, "ANIKA BLESS", 10, 8, 8],
    [12, "ATISEY ELIZABETH", 10, 7, 15],
    [13, "AZORLI AYISHA HAPPY", 10, 9, 10],
    [14, "BAKILA NATASHA",5 , 0, 0],
    [15, "BLAY GETRUDE", 4, 8, 13],
    [16, "BLUDO ISABELLA", 10, 8, 16],
    [17, "DOGBA GENEVIVE", 10, 12, 11],
    [18, "EDEMU ASEYE", 5, 6, 7],
    [19, "ETU GABRIELLA", 10, 14, 15],
    [20, "GAWU ERICA", 4, 9, 13],
    [21, "GORNI EDNA AKUSIKA", 7, 0, 10],
    [22, "KAKRABA PRECIOUS", 10, 10, 16],
    [23, "KOVE GLORIA", 6, 8, 12],
    [24, "KPODO JOY", 10, 14, 18],
    [25, "KPOGO ISABEL", 10, 11, 17],
    [26, "KUMADO PHIDELIA", 6, 6, 13],
    [27, "LETSA PERFECT", 7, 4, 11],
    [28, "LUMORVI HAPPY", 10, 8, 8],
    [29, "NORVIEWU FAITH", 6, 4, 11],
    [30, "RAGNANDA PERPETUAL", 8, 6, 5],
    [31, "RUKAYA ABUBAKAR", 6, 4, 12],
    [32, "TETTEH DEDE JOYCE", 10, 3, 7],
    [33, "TSEY HARRIET", 10, 10, 10],
]

# Create the table
table = doc.add_table(rows=1, cols=len(headers))
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    hdr_cells[i].text = header

# Fill the table
for row in table_data:
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = str(val)
    for _ in range(5):  # Add 5 blank cells for the rest of the columns
        row_cells[len(row)].text = ""

# Save the file
doc.save("SBA_Form_1A_Girls.docx")
print("SBA form created successfully.")
