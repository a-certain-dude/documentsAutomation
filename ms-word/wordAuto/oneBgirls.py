from docx import Document

# Create the Word document
doc = Document()
doc.add_heading('ST. FRANCIS OF ASSISI CATHOLIC JUNIOR HIGH SCHOOL', level=1)
doc.add_heading('SCHOOL BASED ASSESSMENT (SBA) FORM', level=2)
doc.add_paragraph('TERM: ..............     FORM: 1A BOYS     SUBJECT TEACHER: ..............................')
doc.add_paragraph('SUBJECT: .....................................................     YEAR: .......................')

# Define table headers
headers = [
    "NO.", "NAME", "CAT 1\n10", "GROUP WORK\n20", "CAT 2\n20",
    "SUB TOTAL CATS\n50", "END OF TERM EXAMS\n100", "50% OF END OF TERM EXAMS",
    "TOTAL CATS + EXAMS", "GRADE"
]

# Table data from the image
table_data = [
    [1, "ABOTSI SUZZY", 0, 7, 14, 21, 16, 8, 29],
    [2, "ADJETEY STELLA", 0, 2, 15, 17, 16, 8, 25],
    [3, "AFENYO JULIET", 9, 2, 15, 26, 19, 10, 36],
    [4, "AGBESINU PRISCILLA", 6, 2, 14, 22, 20, 10, 32],
    [5, "AGBODZAGAH ESTHER", 6, 4, 15, 25, 20, 10, 35],
    [6, "AHIABU PHILDORS", 0, 1, 13, 14, 24, 12, 26],
    [7, "AKPADZI BLESSING", 0, 0, 8, 8, 8, 4, 12],
    [8, "AMANYO DEBORAH", 9, 5, 14, 28, 24, 12, 40],
    [9, "ATABUU ELIZABETH", 9, 7, 12, 28, 24, 12, 40],
    [10, "ATAKRAH MIRINDA", 10, 6, 14, 30, 20, 10, 40],
    [11, "BADASU VERA", 10, 6, 13, 29, 19, 10, 39],
    [12, "BATIRAM EYRAM", 10, 4, 13, 27, 18, 9, 36],
    [13, "BAYO-VORSAH PHILLIPINE", 10, 4, 15, 29, 20, 10, 39],
    [14, "DEDZO AGNES", 0, 0, 0, 0, 0, 0, 0],  # Add actual scores if available
    [15, "EDZAH ERICA", 0, 7, 8, 15, 16, 8, 23],
    [16, "FIAMOR DELIGHT", 10, 0, 5, 15, 15, 8, 23],
    [17, "GALLEY REJOICE FAFALI", 9, 6, 15, 30, 24, 12, 42],
    [18, "GBADRE RACHAEL", 9, 7, 10, 26, 16, 8, 34],
    [19, "GBEDZE GLORIA", 0, 2, 10, 12, 13, 7, 19],
    [20, "GERALDO FAITH SELASI", 10, 4, 14, 28, 22, 11, 39],
    [21, "KOMLA-DWELL JOAN", 10, 4, 15, 29, 22, 11, 40],
    [22, "KORMEGBE SERAPHINE", 10, 4, 16, 30, 19, 10, 40],
    [23, "KUDEKA FAFALI", 10, 4, 16, 30, 22, 11, 41],
    [24, "KUMAH ABIGAIL", 10, 1, 10, 21, 16, 8, 29],
    [25, "LUMOUR PRECIOUS", 10, 5, 10, 25, 16, 8, 33],
    [26, "MODZAKA BRIDGETTE", 10, 5, 14, 29, 17, 9, 38],
    [27, "QUARSHIGAH ANGELA", 10, 4, 15, 29, 29, 15, 44],
    [28, "SEKLE EWOENAM", 10, 0, 12, 22, 16, 8, 30],
    [29, "SIWAGAH HAPPY", 10, 2, 15, 27, 16, 8, 35],
    [30, "TAY PRISCILLA", 10, 4, 12, 26, 27, 14, 40],
    [31, "TSALIM ZULEIHA", 10, 4, 11, 25, 22, 11, 36],
    [32, "WOTORBU ERICA", 9, 4, 11, 24, 16, 8, 32]
]

# Create the table
table = doc.add_table(rows=1, cols=len(headers))
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, header in enumerate(headers):
    hdr_cells[i].text = header

# Fill the table
for row in table_data:
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = str(val)
    for _ in range(5):  # Add 5 blank cells for the rest of the columns
        row_cells[len(row)].text = ""

# Save the file
doc.save("SBA_Form_1B_Girls.docx")
print("SBA form created successfully.")


